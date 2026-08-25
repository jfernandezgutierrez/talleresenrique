from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r'D:\CLAUDE\taller')
IMG = ROOT / 'DOCUMENTACIÓN' / '03_Revisión_web' / 'capturas_web'
OUT = ROOT / 'DOCUMENTACIÓN' / '03_Revisión_web' / 'Revision_diseno_web_Talleres_Enrique.docx'

BLUE = '1F4D3A'
GREEN = '3E7C59'
LIGHT = 'EAF3ED'
GRAY = '5B6570'

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = sec.bottom_margin = Inches(0.7)
sec.left_margin = sec.right_margin = Inches(0.8)
sec.header_distance = Inches(0.35)
sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1
for st, size, color, before, after in [
    ('Title', 28, BLUE, 0, 10), ('Heading 1', 16, GREEN, 14, 7), ('Heading 2', 13, BLUE, 10, 5)
]:
    s = styles[st]
    s.font.name = 'Calibri'
    s.font.size = Pt(size)
    s.font.color.rgb = RGBColor.from_string(color)
    s.font.bold = True
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn('w:' + m))
        if node is None:
            node = OxmlElement('w:' + m)
            tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'), 'dxa')

def add_textbox(title, lines=3):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.columns[0].width = Inches(6.75)
    cell = t.cell(0, 0); set_cell_margins(cell, 110, 150, 110, 150)
    shade(cell, 'F7F9F8')
    p = cell.paragraphs[0]
    r = p.add_run(title); r.bold = True; r.font.color.rgb = RGBColor.from_string(BLUE)
    for _ in range(lines):
        p = cell.add_paragraph('________________________________________________________________________________')
        p.paragraph_format.space_after = Pt(2)
        p.runs[0].font.color.rgb = RGBColor.from_string('A9B2AD')
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

def add_rating():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run('Valoración general:  '); r.bold = True
    p.add_run('Me gusta [  ]     En parte [  ]     No me gusta [  ]')

def add_questions(items):
    for q in items:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        p.paragraph_format.space_after = Pt(4)
        p.add_run('• ').bold = True
        p.add_run(q)

def add_capture(path, max_h=4.25):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(7)
    p.add_run().add_picture(str(path), width=Inches(6.75), height=None)

def add_header_footer():
    header = sec.header.paragraphs[0]
    header.text = 'REVISIÓN DE DISEÑO WEB · TALLERES ENRIQUE'
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8); header.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
    footer = sec.footer.paragraphs[0]
    footer.text = 'Documento para recoger la opinión del cliente'
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8); footer.runs[0].font.color.rgb = RGBColor.from_string(GRAY)

add_header_footer()

p = doc.add_paragraph(style='Title'); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Revisión del diseño de la web')
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Talleres Enrique'); r.bold = True; r.font.size = Pt(18); r.font.color.rgb = RGBColor.from_string(GREEN)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Cuestionario visual para validar el diseño antes de realizar cambios').italic = True

doc.add_paragraph('\n')
t = doc.add_table(rows=4, cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = False
for row in t.rows:
    row.cells[0].width = Inches(1.65); row.cells[1].width = Inches(5.1)
labels = ['Nombre del cliente', 'Fecha de revisión', 'Persona de contacto', 'Versión revisada']
vals = ['', '', '', 'Web publicada en Netlify · 11/08/2026']
for i, (label, val) in enumerate(zip(labels, vals)):
    c0, c1 = t.rows[i].cells
    shade(c0, LIGHT); set_cell_margins(c0); set_cell_margins(c1)
    c0.paragraphs[0].add_run(label).bold = True
    c1.paragraphs[0].add_run(val if val else '________________________________________')

doc.add_heading('Cómo completar este documento', level=1)
add_questions([
    'Revise cada captura y marque una valoración general.',
    'Indique qué elementos mantendría y cuáles cambiaría.',
    'Si propone un cambio, detalle colores, textos, imágenes o estilo deseado.',
    'Anote al final cualquier comentario que afecte a toda la web.'
])
add_textbox('Primera impresión general de la web', 4)

sections = [
 ('1. Cabecera y portada principal', '01_cabecera_hero.jpg', [
  '¿La cabecera, el logotipo y el menú se entienden con rapidez?',
  '¿Le gustan el color, la imagen de fondo y el estilo general de la portada?',
  '¿El mensaje principal explica bien a qué se dedica el taller?',
  '¿El buscador, WhatsApp y el teléfono tienen la importancia adecuada?',
  '¿Mantendría los selectores de estilo visibles para los visitantes?'
 ]),
 ('2. Sección “Quiénes somos”', '02_quienes_somos.jpg', [
  '¿El texto transmite confianza, cercanía y experiencia?',
  '¿La cantidad de texto es adecuada o debería ser más breve?',
  '¿Añadiría una fotografía real del equipo, del taller o de las instalaciones?',
  '¿El aviso sobre recogida y reparación se entiende y destaca lo suficiente?'
 ]),
 ('3. Servicios · primera parte', '03_servicios_parte_1.jpg', [
  '¿La presentación de los servicios es clara y fácil de recorrer?',
  '¿Los iconos representan bien cada servicio?',
  '¿Cambiaría el orden, el nombre o la descripción de algún servicio?',
  '¿Preferiría fotografías reales en lugar de iconos?'
 ]),
 ('4. Servicios · segunda parte', '04_servicios_parte_2.jpg', [
  '¿Todos los servicios importantes aparecen representados?',
  '¿Las tarjetas tienen suficiente información para el cliente?',
  '¿El botón “Ver todos los servicios” destaca lo necesario?',
  '¿Hay algún servicio que deba añadirse, eliminarse o destacar más?'
 ]),
 ('5. Llamada a la acción', '05_llamada_a_la_accion.jpg', [
  '¿El mensaje invita claramente a contactar?',
  '¿Mantendría los dos botones: WhatsApp y “Ver piezas”?',
  '¿Cambiaría el texto, los colores o el tamaño de esta franja?',
  '¿Añadiría un teléfono o una indicación de horario?'
 ]),
 ('6. Pie de página y datos de contacto', '06_pie_de_pagina.jpg', [
  '¿Los datos de contacto, dirección y horarios son correctos?',
  '¿La navegación y las especialidades incluidas son suficientes?',
  '¿Mantendría Facebook y WhatsApp como accesos destacados?',
  '¿Falta algún dato legal, enlace, red social o ubicación en mapa?'
 ])
]

for idx, (title, filename, questions) in enumerate(sections):
    doc.add_page_break()
    doc.add_heading(title, level=1)
    add_capture(IMG / filename)
    add_rating()
    add_questions(questions)
    add_textbox('¿Qué mantendría de esta parte?', 2)
    add_textbox('¿Qué modificaría y cómo le gustaría verlo?', 3)

doc.add_page_break()
doc.add_heading('Valoración final', level=1)
add_rating()
add_questions([
    '¿El diseño representa correctamente la imagen de Talleres Enrique?',
    '¿La web transmite profesionalidad y confianza?',
    '¿La información más importante se encuentra con facilidad?',
    '¿El estilo es adecuado para los clientes habituales del negocio?',
    '¿Autoriza mantener este diseño como base para la versión definitiva?  Sí [  ]  No [  ]'
])
add_textbox('Tres cambios prioritarios', 5)
add_textbox('Otros comentarios o referencias de webs que le gusten', 5)

doc.core_properties.title = 'Revisión del diseño web - Talleres Enrique'
doc.core_properties.subject = 'Cuestionario visual de validación con cliente'
doc.core_properties.author = 'Talleres Enrique'
doc.save(OUT)
print(OUT)
