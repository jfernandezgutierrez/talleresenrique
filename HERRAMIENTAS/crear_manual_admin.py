from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
CAP = ROOT / 'DOCUMENTACIÓN' / '05_Manual_admin' / 'capturas'
OUT = ROOT / 'DOCUMENTACIÓN' / '05_Manual_admin' / 'Manual_de_usuario_Talleres_Enrique_Admin_Nitido.docx'
GREEN = RGBColor(30, 91, 60)
LIGHT_GREEN = 'E8F2EC'
YELLOW = RGBColor(214, 164, 0)
GRAY = RGBColor(90, 98, 105)

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(11); sec.page_height = Inches(8.5)
sec.top_margin = Inches(0.58); sec.bottom_margin = Inches(0.58)
sec.left_margin = Inches(0.65); sec.right_margin = Inches(0.65)

styles = doc.styles
normal = styles['Normal']; normal.font.name = 'Aptos'; normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(5); normal.paragraph_format.line_spacing = 1.12
for name, size, color in [('Title', 30, GREEN), ('Heading 1', 21, GREEN), ('Heading 2', 15, GREEN), ('Heading 3', 11.5, YELLOW)]:
    st = styles[name]; st.font.name = 'Aptos Display'; st.font.size = Pt(size); st.font.bold = True; st.font.color.rgb = color
    st.paragraph_format.space_before = Pt(10); st.paragraph_format.space_after = Pt(5); st.paragraph_format.keep_with_next = True

caption = styles.add_style('Caption Manual', WD_STYLE_TYPE.PARAGRAPH)
caption.font.name='Aptos'; caption.font.size=Pt(8.5); caption.font.italic=True; caption.font.color.rgb=GRAY
caption.paragraph_format.space_before=Pt(3); caption.paragraph_format.space_after=Pt(8)

def shade(paragraph, fill=LIGHT_GREEN):
    pPr = paragraph._p.get_or_add_pPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); pPr.append(shd)

def callout(label, text, warning=False):
    p=doc.add_paragraph(); shade(p, 'FFF4CE' if warning else LIGHT_GREEN)
    p.paragraph_format.left_indent=Inches(.14); p.paragraph_format.right_indent=Inches(.14)
    p.paragraph_format.space_before=Pt(5); p.paragraph_format.space_after=Pt(7)
    r=p.add_run(label+' '); r.bold=True; r.font.color.rgb=GREEN if not warning else RGBColor(120,75,0)
    p.add_run(text)

def bullet(text): doc.add_paragraph(text, style='List Bullet')
def step(text): doc.add_paragraph(text, style='List Number')
def image(name, caption_text, width=9.55):
    path=CAP/name
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.keep_with_next=True
    p.add_run().add_picture(str(path), width=Inches(width))
    c=doc.add_paragraph(caption_text, style='Caption Manual'); c.alignment=WD_ALIGN_PARAGRAPH.CENTER

def new_section(title, intro=None):
    doc.add_page_break(); doc.add_heading(title, 1)
    if intro: doc.add_paragraph(intro)

# Cover: editorial_cover
doc.add_paragraph('TALLERES ENRIQUE', style='Subtitle').alignment=WD_ALIGN_PARAGRAPH.CENTER
p=doc.add_paragraph(style='Title'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(110)
p.add_run('Manual de usuario')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Panel de administración'); r.font.size=Pt(18); r.bold=True; r.font.color.rgb=GREEN
p=doc.add_paragraph('Guía completa para gestionar la web sin conocimientos técnicos')
p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].font.size=Pt(12); p.runs[0].font.color.rgb=GRAY
p=doc.add_paragraph('Versión 1.0 · Agosto de 2026'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before=Pt(80); p.runs[0].font.bold=True; p.runs[0].font.color.rgb=YELLOW
callout('Objetivo.', 'Explicar, paso a paso, cómo actualizar piezas, categorías, servicios, páginas y el contenido general de la web.')

new_section('1. Antes de empezar')
doc.add_paragraph('El panel permite cambiar el contenido público de Talleres Enrique. Los cambios guardados se almacenan en la base de datos y pueden aparecer inmediatamente en la web.')
doc.add_heading('Acceso', 2)
step('Abre https://tallerenriqueadmin.netlify.app/.')
step('Introduce el correo y la contraseña facilitados por el responsable del proyecto.')
step('Pulsa “Entrar al panel”.')
callout('Seguridad.', 'No compartas la contraseña, no la guardes en equipos públicos y pulsa “Cerrar sesión” al terminar.', True)
doc.add_heading('Reglas sencillas para trabajar con seguridad', 2)
bullet('Pulsa Guardar únicamente cuando hayas revisado el contenido.')
bullet('Usa Actualizar para volver a cargar los datos del servidor.')
bullet('Antes de eliminar algo, comprueba que no está siendo utilizado.')
bullet('Después de un cambio importante, abre la web pública y comprueba el resultado.')
doc.add_heading('Menú lateral', 2)
doc.add_paragraph('PZ = Piezas · CT = Categorías · SV = Servicios · PG = Páginas · CF = Configuración. El botón ‹ reduce el menú; ↪ cierra la sesión.')

new_section('2. Gestión de piezas', 'Este apartado controla los productos y recambios del catálogo público.')
image('01-piezas.png', 'Pantalla principal de Piezas: filtros, resumen de stock, tabla y paginación.')
doc.add_heading('Buscar y ordenar', 2)
bullet('Escribe un nombre o referencia en el buscador.')
bullet('Filtra por categoría o por estado de stock.')
bullet('Pulsa Nombre, Referencia, Categoría o Stock para cambiar el orden.')
bullet('Usa los números de página para recorrer todo el catálogo.')
doc.add_heading('Crear una pieza', 2)
image('02-nueva-pieza.png', 'Formulario “Nueva pieza”. Los campos con asterisco son obligatorios.')
step('Pulsa “+ Nueva pieza”.')
step('Escribe el nombre y la referencia. Elige una categoría.')
step('Selecciona un icono y completa compatibilidad y descripción.')
step('Elige el estado de stock y revisa la etiqueta que verá el cliente.')
step('Añade imágenes arrastrándolas o seleccionándolas. También puedes añadir una URL de YouTube/Vimeo o subir un vídeo.')
step('Pulsa “Añadir pieza”.')
callout('Imágenes.', 'Formatos JPG, PNG o WEBP; máximo 5 MB por imagen. El vídeo admite hasta 100 MB. Marca como principal la foto que debe verse primero.')
doc.add_heading('Editar o eliminar', 2)
doc.add_paragraph('Editar abre el mismo formulario con los datos actuales. Eliminar solicita confirmación y no se puede deshacer. Si solo falta stock, conviene cambiar el estado a “Sin stock / Consultar” en vez de borrar la pieza.')

new_section('3. Categorías', 'Las categorías organizan las piezas y alimentan los filtros del catálogo.')
image('03-categorias.png', 'Listado de categorías y número de piezas asignadas a cada una.')
image('04-nueva-categoria.png', 'Alta de categoría: nombre, icono, color y vista previa.')
step('Pulsa “+ Nueva categoría”.')
step('Indica un nombre corto y comprensible.')
step('Escoge un icono y un color que permitan reconocerla rápidamente.')
step('Pulsa “Crear categoría”.')
callout('Antes de eliminar.', 'Comprueba el contador de piezas. Una categoría utilizada puede dejar productos mal clasificados; reasigna primero sus piezas.', True)

new_section('4. Servicios', 'Controla las tarjetas de la página Servicios y su presentación en la portada.')
image('05-servicios.png', 'Servicios configurados, número de prestaciones y orden de aparición.')
image('06-nuevo-servicio.png', 'Formulario para crear o editar un servicio.')
step('Pulsa “+ Nuevo servicio”.')
step('Completa título, icono y descripción.')
step('Escribe una prestación por línea; cada línea aparecerá con una marca de verificación.')
step('Asigna el orden. Los números menores aparecen antes; se recomienda usar 10, 20, 30…')
step('Activa “Visible en la web pública” o desactívalo para ocultarlo sin borrarlo.')
step('Pulsa “Guardar servicio”.')

new_section('5. Visibilidad de páginas', 'Permite ocultar temporalmente partes completas de la web.')
image('07-paginas.png', 'Interruptores de Servicios, Catálogo y Contacto.')
doc.add_paragraph('Desmarca una página y pulsa “Guardar visibilidad”. Desaparecerá del menú y quien visite su dirección será enviado a Inicio. Para recuperarla, vuelve a marcarla y guarda.')
callout('Siempre disponibles.', 'Inicio y las páginas legales no se pueden desactivar desde aquí.')

new_section('6. Configuración general', 'Aquí se editan los datos del negocio y la mayoría de los textos públicos.')
image('08-configuracion-contacto.png', 'Contacto, horarios, presentación e indicadores de portada.')
doc.add_heading('01 · Contacto', 2)
bullet('Teléfono visible: formato fácil de leer. Enlace tel: solo números y prefijo internacional, por ejemplo +34942590301.')
bullet('WhatsApp: número con prefijo de país y sin espacios ni símbolo +.')
bullet('Email, Facebook y dirección se muestran en Contacto y/o pie de página.')
doc.add_heading('02 · Horario y respuesta', 2)
doc.add_paragraph('Actualiza los horarios tal como deben leerlos los clientes. El plazo de respuesta acompaña a los botones de WhatsApp.')
doc.add_heading('03–04 · Portada e indicadores', 2)
doc.add_paragraph('El título, subtítulo, mensaje, “Quiénes somos”, política de recogida y las tres cifras destacadas aparecen en Inicio. Evita afirmaciones o cifras que no puedan justificarse.')

new_section('7. Portada, marcas y ventajas')
image('09-configuracion-portada.png', 'Portada completa, marcas y bloque “Por qué elegirnos”.')
doc.add_heading('05 · Portada completa', 2)
bullet('La imagen puede indicarse mediante URL o subirse desde el ordenador.')
bullet('El distintivo superior es una frase breve, por ejemplo “Distribuidores Rapid · Toda Cantabria”.')
bullet('Escribe una marca por línea. El orden de las líneas será el orden visible.')
bullet('El título y descripción del bloque de servicios presentan las tarjetas en Inicio.')
doc.add_heading('06 · Por qué elegirnos', 2)
doc.add_paragraph('Cada ventaja ocupa una línea con el formato icono|título|descripción. Mantén exactamente los dos separadores verticales. Ejemplo:')
callout('Ejemplo.', 'trophy|+20 años de experiencia|Décadas de trabajo con maquinaria agrícola en Cantabria.')

new_section('8. Títulos, llamadas a la acción y textos legales')
image('10-configuracion-textos.png', 'Títulos y descripciones de páginas y llamadas a la acción.')
doc.add_heading('07 · Títulos y CTA', 2)
doc.add_paragraph('Puedes cambiar los encabezados visibles de Servicios, Catálogo y Contacto, además de los mensajes que invitan al cliente a escribir por WhatsApp. Usa títulos breves y descripciones claras.')
image('11-configuracion-legales.png', 'Edición de aviso legal, privacidad y cookies.')
doc.add_heading('08 · Textos legales', 2)
bullet('Actualiza la fecha de revisión cuando se modifique cualquiera de los textos.')
bullet('Mantén marcado “Mostrar aviso pendiente de validar” hasta que los datos fiscales y textos hayan sido revisados.')
bullet('Separa párrafos con una línea en blanco.')
callout('Importante.', 'Los textos legales deben ser validados por el titular y, cuando proceda, por un profesional.', True)

new_section('9. WhatsApp y banner temporal')
image('12-configuracion-whatsapp-banner.png', 'Mensajes automáticos de WhatsApp y banner con programación opcional.')
doc.add_heading('09 · Mensajes de WhatsApp', 2)
doc.add_paragraph('Estos textos aparecen preparados cuando un visitante pulsa un botón de WhatsApp. En la consulta de piezas conserva {pieza} y {referencia}; el sistema las sustituye automáticamente por el producto elegido.')
doc.add_heading('10 · Banner temporal', 2)
step('Marca “Activar banner”.')
step('Escribe título y mensaje, y selecciona Información, Aviso o Destacado.')
step('Opcionalmente indica inicio y fin para vacaciones, cierres u horarios especiales.')
step('Añade un texto y una URL si el banner debe incluir un enlace.')
step('Pulsa “Guardar configuración”.')
callout('Fechas vacías.', 'Si no introduces inicio ni fin, el banner permanecerá visible mientras esté activado.')

new_section('10. Guardado, comprobación y resolución de problemas')
doc.add_heading('Después de cada cambio', 2)
step('Pulsa el botón Guardar correspondiente.')
step('Espera a ver “Cambios guardados”.')
step('Abre o actualiza la web pública y comprueba el resultado en ordenador y móvil.')
doc.add_heading('Si un cambio no aparece', 2)
bullet('Pulsa “Actualizar” en el panel.')
bullet('Recarga la web pública con Ctrl+F5.')
bullet('Comprueba que el servicio o la página están visibles.')
bullet('Verifica que has pulsado Guardar y que no apareció un error de conexión.')
doc.add_heading('Si una imagen no carga', 2)
bullet('Comprueba que el formato sea JPG, PNG o WEBP y que no exceda el límite.')
bullet('Si usas una URL, debe comenzar por https:// y ser accesible públicamente.')
doc.add_heading('Errores que conviene evitar', 2)
bullet('Eliminar categorías con piezas asociadas.')
bullet('Borrar productos cuando basta con cambiar su stock.')
bullet('Quitar las variables {pieza} y {referencia} del mensaje de catálogo.')
bullet('Publicar textos legales sin validación.')
bullet('Cerrar el navegador antes de que termine una subida.')

new_section('11. Lista rápida de comprobación')
for text in ['Datos de contacto y horarios correctos','Páginas necesarias activadas','Piezas con categoría, stock e imagen principal','Servicios visibles y ordenados','Textos de portada revisados','WhatsApp abre con el mensaje esperado','Banner desactivado o con fechas correctas','Textos legales validados','Web pública comprobada después de guardar','Sesión cerrada al terminar']:
    bullet(text)
callout('Soporte.', 'Si aparece un error persistente, anota qué estabas haciendo, copia el mensaje y realiza una captura antes de contactar con la persona responsable del mantenimiento.')

# Header/footer
for section in doc.sections:
    h=section.header.paragraphs[0]; h.text='Talleres Enrique · Manual del panel de administración'; h.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    h.runs[0].font.size=Pt(8); h.runs[0].font.color.rgb=GRAY
    f=section.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.CENTER
    f.add_run('Manual de usuario · Agosto de 2026').font.size=Pt(8)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
